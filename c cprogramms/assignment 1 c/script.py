import os
import re
from pygments import highlight
from pygments.lexers import CLexer
from pygments.formatters import ImageFormatter
from PIL import Image
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from docx import Document

# -------------------------
# Natural sort key (so q2.c comes before q10.c)
# -------------------------
def natural_sort_key(text):
    return [int(s) if s.isdigit() else s.lower() for s in re.split(r'(\d+)', text)]

# -------------------------
# 1. Collect all .c files in natural order
# -------------------------
code_files = [f for f in os.listdir() if f.endswith(".c")]
code_files.sort(key=natural_sort_key)

if not code_files:
    print("‚ö†Ô∏è No .c files found in this directory.")
    exit()

# -------------------------
# 2. Create screenshots of each code file (dark mode)
# -------------------------
images = []
for i, filename in enumerate(code_files, start=1):
    with open(filename, "r", encoding="utf-8") as f:
        code = f.read()

    formatter = ImageFormatter(
        font_name="Consolas",  # safe on Windows, fallback to DejaVu Sans Mono on Linux
        font_size=18,
        line_numbers=True,
        style="monokai"  # dark mode theme
    )
    img_data = highlight(code, CLexer(), formatter)

    img_name = f"question{i}.png"
    with open(img_name, "wb") as img_file:
        img_file.write(img_data)

    images.append((i, img_name))
    print(f"‚úÖ Created screenshot: {img_name} from {filename}")

# -------------------------
# 3. Make PDF from images
# -------------------------
pdf_file = "Questions.pdf"
c = canvas.Canvas(pdf_file, pagesize=A4)
width, height = A4

for i, (q_num, img_name) in enumerate(images, start=1):
    # Heading
    c.setFont("Helvetica-Bold", 18)
    c.drawString(200, height - 50, f"Question {q_num}")

    try:
        img = Image.open(img_name)
        aspect = img.width / img.height

        new_width = width - 80
        new_height = new_width / aspect
        if new_height > height - 150:
            new_height = height - 150
            new_width = new_height * aspect

        c.drawImage(img_name, 40, height - new_height - 100, new_width, new_height)
    except Exception as e:
        print(f"‚ö†Ô∏è Could not add {img_name}: {e}")

    c.showPage()

c.save()
print("üìÑ PDF created:", pdf_file)

# -------------------------
# 4. Make Word file with code text
# -------------------------
doc = Document()
for i, filename in enumerate(code_files, start=1):
    doc.add_heading(f"Question {i}", level=1)
    with open(filename, "r", encoding="utf-8") as f:
        code = f.read()
    doc.add_paragraph(code, style="Normal")

word_file = "Questions.docx"
doc.save(word_file)
print("üìÑ Word file created:", word_file)
