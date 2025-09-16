import os
import subprocess
import re
from pygments import highlight
from pygments.lexers import PythonLexer
from pygments.formatters import ImageFormatter
from PIL import Image
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Preformatted
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document

# -------------------------
# Natural sort (Q1, Q2, Q10 correctly ordered)
# -------------------------
def natural_sort_key(s):
    return [int(text) if text.isdigit() else text.lower()
            for text in re.split(r'(\d+)', s)]

# -------------------------
# 1. Collect all Q*.py files in natural order
# -------------------------
code_files = [f for f in os.listdir() if f.startswith("Q") and f.endswith(".py") and f != "compile_questions.py"]
code_files.sort(key=natural_sort_key)

# -------------------------
# 2. Create screenshots of each code file
# -------------------------
images = []
for filename in code_files:
    with open(filename, "r") as f:
        code = f.read()

    formatter = ImageFormatter(
        font_name="Consolas",
        font_size=18,
        line_numbers=True,
        style="monokai"
    )
    img_data = highlight(code, PythonLexer(), formatter)

    img_name = f"{filename.replace('.py', '')}.png"
    with open(img_name, "wb") as img_file:
        img_file.write(img_data)

    images.append((filename, img_name))
    print(f"✅ Created screenshot: {img_name} from {filename}")

# -------------------------
# 3. PDF with screenshots only
# -------------------------
pdf_code_file = "All_Questions_Code.pdf"
c = canvas.Canvas(pdf_code_file, pagesize=A4)
width, height = A4

for filename, img_name in images:
    c.setFont("Helvetica-Bold", 18)
    c.drawString(160, height - 50, filename)

    try:
        img = Image.open(img_name)
        aspect = img.width / img.height

        new_width = width - 80
        new_height = new_width / aspect
        if new_height > height - 200:
            new_height = height - 200
            new_width = new_height * aspect

        c.drawImage(img_name, 40, height - new_height - 100, new_width, new_height)
    except Exception as e:
        print(f"⚠️ Could not add {img_name}: {e}")

    c.showPage()

c.save()
print("📄 PDF created:", pdf_code_file)

# -------------------------
# 4. PDF with compiled code + outputs as text
# -------------------------
pdf_file = "All_Questions.pdf"
doc_pdf = SimpleDocTemplate(pdf_file, pagesize=A4)
styles = getSampleStyleSheet()
story = []

for filename in code_files:
    story.append(Paragraph(f"<b>{filename}</b>", styles['Title']))
    story.append(Spacer(1, 12))

    # Add code
    with open(filename, "r") as f:
        code = f.read()
    story.append(Paragraph("<b>Code:</b>", styles['Heading3']))
    story.append(Preformatted(code, styles['Code']))
    story.append(Spacer(1, 12))

    # Run code and capture output
    output_text = ""
    try:
        result = subprocess.run(["python", filename], capture_output=True, text=True, timeout=5)
        if result.returncode == 0:
            output_text = result.stdout
    except Exception:
        output_text = ""

    if output_text.strip():
        story.append(Paragraph("<b>Output:</b>", styles['Heading3']))
        story.append(Preformatted(output_text, styles['Code']))
        story.append(Spacer(1, 24))

doc_pdf.build(story)
print("📄 PDF created:", pdf_file)

# -------------------------
# 5. Word file with code + outputs
# -------------------------
word_file = "All_Questions.docx"
doc_word = Document()

for filename in code_files:
    doc_word.add_heading(filename, level=1)

    with open(filename, "r") as f:
        code = f.read()
    doc_word.add_paragraph(code, style="Normal")

    output_text = ""
    try:
        result = subprocess.run(["python", filename], capture_output=True, text=True, timeout=5)
        if result.returncode == 0:
            output_text = result.stdout
    except Exception:
        output_text = ""

    if output_text.strip():
        doc_word.add_heading("Output:", level=2)
        doc_word.add_paragraph(output_text, style="Normal")

doc_word.save(word_file)
print("📄 Word file created:", word_file)
