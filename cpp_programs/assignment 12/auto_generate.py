import os
from pygments import highlight
from pygments.lexers import CppLexer
from pygments.formatters import ImageFormatter
from PIL import Image
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from docx import Document

# -------------------------
# 1. Collect all .cpp files in alphabetical order
# -------------------------
code_files = [f for f in os.listdir() if f.endswith(".cpp")]
code_files.sort()  # alphabetical order

# -------------------------
# 2. Create screenshots of each code file (dark mode)
# -------------------------
images = []
for i, filename in enumerate(code_files, start=1):
    with open(filename, "r") as f:
        code = f.read()

    formatter = ImageFormatter(
        font_name="Consolas",  # safe on Windows
        font_size=18,
        line_numbers=True,
        style="monokai"  # dark mode theme
    )
    img_data = highlight(code, CppLexer(), formatter)

    img_name = f"case{i}.png"
    with open(img_name, "wb") as img_file:
        img_file.write(img_data)

    images.append((i, img_name))
    print(f"✅ Created screenshot: {img_name} from {filename}")

# -------------------------
# 3. Make PDF from images
# -------------------------
pdf_file = "Case_Studies.pdf"
c = canvas.Canvas(pdf_file, pagesize=A4)
width, height = A4

for i, (case_num, img_name) in enumerate(images, start=1):
    # Heading
    c.setFont("Helvetica-Bold", 18)
    c.drawString(200, height - 50, f"Case Study {case_num}")

    try:
        img = Image.open(img_name)
        aspect = img.width / img.height

        new_width = width - 80
        new_height = new_width / aspect
        if new_height > height - 150:
            new_height = height - 150
            new_width = new_height * aspect

        c.drawImage(img_name, 40, height - new_height - 100, new_width, new_height)
    except Exception as e:
        print(f"⚠️ Could not add {img_name}: {e}")

    c.showPage()

c.save()
print("📄 PDF created:", pdf_file)

# -------------------------
# 4. Make Word file with code text
# -------------------------
doc = Document()
for i, filename in enumerate(code_files, start=1):
    doc.add_heading(f"Case Study {i}", level=1)
    with open(filename, "r") as f:
        code = f.read()
    doc.add_paragraph(code, style="Normal")

word_file = "Case_Studies.docx"
doc.save(word_file)
print("📄 Word file created:", word_file)
